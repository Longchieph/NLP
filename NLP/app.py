import os
import torch
import nltk
from flask import Flask, render_template, request, send_file, redirect, url_for
from werkzeug.utils import secure_filename
from transformers import MarianMTModel, MarianTokenizer
from docx import Document
import fitz  # PyMuPDF
from pdfminer.high_level import extract_text
import re
import html
import ftfy
import unicodedata
nltk.download('punkt')

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config['JSON_AS_ASCII'] = False
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024
SUPPORTED_PAIRS = ['en-vi', 'vi-en', 'en-fr', 'fr-en','vi-fr', 'fr-vi']
models_cache = {}

def get_model(src_lang, dest_lang):
    key = f"{src_lang}-{dest_lang}"
    if key not in SUPPORTED_PAIRS:
        raise ValueError(f"Unsupported language pair: {key}. Supported pairs: {', '.join(SUPPORTED_PAIRS)}")
    
    if key not in models_cache:
        model_path = f"Helsinki-NLP/opus-mt-{key}" if key not in ["en-vi", "vi-en"] else f"./results/fine-tuned-{key}"
        tokenizer = MarianTokenizer.from_pretrained(model_path)
        model = MarianMTModel.from_pretrained(model_path)
        models_cache[key] = (model, tokenizer)
    
    return models_cache[key]

def split_sentences(text):
    return nltk.sent_tokenize(text)


def normalize_text(text):
    # Sửa lỗi mã hóa và chuẩn hóa Unicode
    text = ftfy.fix_text(text)
    text = html.unescape(text)
    text = unicodedata.normalize('NFC', text)
    # Loại bỏ các ký tự control không cần thiết
    text = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', text)
    return re.sub(r'\s+', ' ', text).strip()
def clean_text(text):
    return normalize_text(text)

def translate_text(text, model, tokenizer):
    sentences = split_sentences(text)
    translated = []
    
    for sent in sentences:
        inputs = tokenizer(sent, return_tensors="pt", truncation=True, max_length=512)
        with torch.no_grad():
            outputs = model.generate(**inputs)
        translated.append(tokenizer.decode(outputs[0], skip_special_tokens=True))
    
    return clean_text(" ".join(translated))

def translate_word_file(filepath, model, tokenizer):
    doc = Document(filepath)
    
    # Translate paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraph.text = translate_text(paragraph.text, model, tokenizer)
    
    # Translate tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cell.text = translate_text(cell.text, model, tokenizer)
    
    # Translate headers/footers
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                if paragraph.text.strip():
                    paragraph.text = translate_text(paragraph.text, model, tokenizer)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                if paragraph.text.strip():
                    paragraph.text = translate_text(paragraph.text, model, tokenizer)
    
    translated_path = os.path.join(app.config["UPLOAD_FOLDER"], f"translated_{os.path.basename(filepath)}")
    doc.save(translated_path)
    return translated_path

def translate_pdf_file(filepath, model, tokenizer):
    try:
        with fitz.open(filepath) as doc:
            translated_doc = fitz.open()
            font_path = os.path.join("templates", "ARIAL.TTF")
            # Sử dụng font hỗ trợ tiếng Việt
            try:
                
                font = fitz.Font(fontfile=font_path)
                font_name = font.name.replace(" ", "")  # Lấy tên font
            except:
                font_name = "times"  # Fallback font
            
            for page in doc:
                # Tạo trang mới với kích thước giống nguyên bản
                new_page = translated_doc.new_page(
                    width=page.rect.width,
                    height=page.rect.height
                )
                
                # Sao chép hình ảnh từ trang gốc
                for img in page.get_images(full=True):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    new_page.insert_image(
                        rect=page.get_image_bbox(img),
                        stream=base_image["image"],
                        keep_proportion=True
                    )
                
                # Xử lý văn bản
                for block in page.get_text("dict")["blocks"]:
                    if block["type"] == 0:  # Chỉ xử lý khối văn bản
                        for line in block["lines"]:
                            for span in line["spans"]:
                                # Chuẩn hóa văn bản
                                raw_text = normalize_text(span["text"])
                                translated = translate_text(raw_text, model, tokenizer)
                                
                                # Xử lý màu sắc
                                try:
                                    color = tuple([c/255 for c in span["color"][:3]])  # Chuẩn hóa màu sắc
                                except:
                                    color = (0, 0, 0)  # Mặc định là màu đen
                                
                                # Chèn văn bản đã dịch vào trang mới
                                new_page.insert_text(
                                    point=span["origin"],  # Giữ nguyên vị trí
                                    text=translated,
                                    fontsize=span["size"],  # Giữ nguyên kích thước font
                                    fontname=font_name,  # Sử dụng tên font
                                    color=color,
                                    render_mode=0  # Chế độ render mặc định
                                )
            
            # Lưu file PDF mới
            translated_filename = f"translated_{os.path.basename(filepath)}"
            translated_path = os.path.join(app.config['UPLOAD_FOLDER'], translated_filename)
            translated_doc.save(translated_path, garbage=4, deflate=True, clean=True)
            
            return translated_path
    
    except Exception as e:
        print(f"PDF Error: {str(e)}")
        raise RuntimeError(f"PDF processing failed: {str(e)}")
ALLOWED_EXTENSIONS = {'docx', 'pdf'}

@app.route('/', methods=['GET', 'POST'])
def translate():
    translated_text = None
    error_msg = None
    translated_file = None
    src_lang = request.form.get('src_lang', 'en')
    dest_lang = request.form.get('dest_lang', 'vi')
    languages = [{"code": "en", "name": "English"}, {"code": "vi", "name": "Vietnamese"}]

    if request.method == 'POST':
        try:
            if not src_lang or not dest_lang:
                raise ValueError("Please select both source and target languages")
            
            model, tokenizer = get_model(src_lang, dest_lang)

            if 'text' in request.form and request.form['text'].strip():
                translated_text = translate_text(request.form['text'], model, tokenizer)

            if 'file' in request.files:
                file = request.files['file']
                if file and allowed_file(file.filename):
                    filename = secure_filename(file.filename)
                    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
                    file.save(filepath)
                    
                    if filename.endswith('.docx'):
                        translated_file = translate_word_file(filepath, model, tokenizer)
                    elif filename.endswith('.pdf'):
                        translated_file = translate_pdf_file(filepath, model, tokenizer)
                    
                    return redirect(url_for('download_translated_file', filename=os.path.basename(translated_file)))
                else:
                    raise ValueError("Invalid file format. Supported formats: .docx, .pdf")

        except Exception as e:
            error_msg = str(e)
    
    return render_template('index.html', 
                         translated_text=translated_text,
                         error_msg=error_msg,
                         src_lang=src_lang,
                         dest_lang=dest_lang,
                         languages=languages)

@app.route('/download/<filename>')
def download_translated_file(filename):
    return send_file(os.path.join(app.config["UPLOAD_FOLDER"], filename), as_attachment=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

if __name__ == '__main__':
    app.run(debug=True)